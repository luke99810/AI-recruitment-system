"""
文档解析模块：支持 PDF / Word / TXT 格式的 JD 和简历解析。
将非结构化文档转换为纯文本，供后续 LLM 处理。

PDF 解析采用三级降级策略：
  1. PyMuPDF 提取文本层（适用于文字型 PDF）
  2. pdfplumber 兜底提取（处理 PyMuPDF 提取失败的文字型 PDF）
  3. RapidOCR 识别图片（适用于扫描件/图片型 PDF）
"""
import io
import logging
from pathlib import Path
from typing import Optional

# ── onnxruntime 必须在 fitz 之前载入，顺序不能反 ──
#
# PyMuPDF 一被 import 就把系统的 msvcp140.dll 拉进进程。DLL 在同一进程内
# 按名字唯一、谁先载入谁生效，之后 onnxruntime 只能用到那一份。
# 某些 Windows 上 System32 的 msvcp140.dll 停在旧版（实测 14.00.24215.1，
# 而同机 vcruntime140_1.dll 已是 14.50 —— 被别的安装程序覆盖回去了），
# onnxruntime 需要 ≥14.4x，于是抛
#   ImportError: DLL load failed while importing onnxruntime_pybind11_state
# 整条 OCR 降级路径随之失效：扫描件 PDF 一律解析不出文字。
#
# 实测：先 import fitz 再 import onnxruntime → 失败；反过来 → 成功。
# pdfplumber / cv2 / numpy 都不触发，只有 fitz。
#
# 没装 OCR 也不影响：import 失败就置 None，第 3 级自然降级。
try:
    import onnxruntime as _onnxruntime  # noqa: F401
    _OCR_RUNTIME_ERROR: Optional[str] = None
except Exception as _e:  # pragma: no cover - 取决于宿主环境
    _onnxruntime = None
    _OCR_RUNTIME_ERROR = str(_e)

import fitz  # PyMuPDF  # noqa: E402

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """解析 PDF 文件，三级降级提取全部文本。"""
    # ---- 第 1 级：PyMuPDF ----
    text = _extract_with_pymupdf(file_bytes)
    if text:
        logger.info("[parser] PDF 文本由 PyMuPDF 提取，长度=%d", len(text))
        return text

    # ---- 第 2 级：pdfplumber ----
    text = _extract_with_pdfplumber(file_bytes)
    if text:
        logger.info("[parser] PDF 文本由 pdfplumber 提取，长度=%d", len(text))
        return text

    # ---- 第 3 级：OCR ----
    text = _extract_with_ocr(file_bytes)
    if text:
        logger.info("[parser] PDF 文本由 OCR 提取，长度=%d", len(text))
        return text

    return ""


def _extract_with_pymupdf(file_bytes: bytes) -> str:
    """使用 PyMuPDF 提取文本层。"""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            t = page.get_text("text")
            if t:
                text_parts.append(t)
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning("[parser] PyMuPDF 提取失败: %s", e)
        return ""


def _extract_with_pdfplumber(file_bytes: bytes) -> str:
    """使用 pdfplumber 兜底提取文本。"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning("[parser] pdfplumber 提取失败: %s", e)
        return ""


def ocr_availability() -> tuple[bool, str]:
    """OCR 这一级到底能不能用，以及不能用的原因。

    区分三种情况，因为它们的处置完全不同：
      - 没装依赖        → `pip install rapidocr-onnxruntime`
      - 装了但环境坏     → 宿主机运行库问题（见文件顶部注释），装包解决不了
      - 可用            → 识别不出文字才是"这份 PDF 不行"
    只报"OCR 未能识别文字"会把后两种混为一谈，用户按提示换文件也没用。
    """
    if _onnxruntime is None:
        return False, f"onnxruntime 无法载入：{_OCR_RUNTIME_ERROR}"
    try:
        import rapidocr_onnxruntime  # noqa: F401
    except ImportError:
        return False, "未安装 rapidocr-onnxruntime（可选依赖）"
    return True, ""


def _extract_with_ocr(file_bytes: bytes) -> str:
    """使用 RapidOCR 对扫描件/图片型 PDF 做 OCR 识别。"""
    ok, reason = ocr_availability()
    if not ok:
        logger.warning("[parser] OCR 不可用，跳过第 3 级：%s", reason)
        return ""
    try:
        from rapidocr_onnxruntime import RapidOCR
        ocr = RapidOCR()

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")
            result, _ = ocr(img_bytes)
            if result:
                page_text = "\n".join(line[1] for line in result if line and len(line) > 1)
                if page_text:
                    text_parts.append(page_text)
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning("[parser] OCR 提取失败: %s", e)
        return ""


def parse_docx(file_bytes: bytes) -> str:
    """解析 Word (.docx) 文件，提取段落和表格文本。"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts).strip()


def parse_txt(file_bytes: bytes) -> str:
    """解析纯文本文件，自动处理编码。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except (UnicodeDecodeError, ValueError):
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    根据文件扩展名自动选择解析器。

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名（用于判断格式）

    Returns:
        提取的纯文本

    Raises:
        ValueError: 不支持的文件格式 或 解析后内容为空
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text = parse_docx(file_bytes)
    elif ext in (".txt", ".md"):
        text = parse_txt(file_bytes)
    else:
        # 尝试根据内容嗅探
        if file_bytes[:4] == b"%PDF":
            text = parse_pdf(file_bytes)
        else:
            text = parse_txt(file_bytes)

    if not text:
        ok, reason = ocr_availability()
        if ext == ".pdf" and not ok:
            raise ValueError(
                f"文件 '{filename}' 没有文本层（图片型/扫描件 PDF），"
                f"而 OCR 当前不可用：{reason}。"
                f"请改用文字型 PDF 或 .docx/.txt，或先修好 OCR 依赖。"
            )
        raise ValueError(
            f"文件 '{filename}' 解析后内容为空。"
            f"该文件可能是图片型/扫描件 PDF，OCR 已运行但未识别出文字，"
            f"请尝试上传文字型 PDF 或 .docx/.txt 格式简历。"
        )

    return text


def parse_uploaded_file(file_content: bytes, filename: str) -> str:
    """
    对外统一接口：解析上传的文件，返回纯文本。
    处理流程：文件字节 → 格式判断 → 对应解析器 → 纯文本
    """
    try:
        return parse_document(file_content, filename)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"解析文件 '{filename}' 失败: {str(e)}")


# 场景B兼容别名
parse_file = parse_uploaded_file
